# all the imports
# DOCX stuff
# https://automatetheboringstuff.com/chapter13/
# https://media.readthedocs.org/pdf/python-docx/latest/python-docx.pdf
# https://python-docx.readthedocs.io/en/latest/genindex.html

import docx

def docx_demo(wk_path):
# Part 1
#new empty Document - uses default template
	doc_work = docx.Document()
#Add headers lines (paragraphs)
	doc_work.add_heading('This is a Main Heading', 0)
	doc_work.add_heading('This is a level 1 Heading', 1)
	doc_work.add_heading('This is a level 2 Heading', 2)
	doc_work.add_heading('This is a level 3 Heading', 3)
	doc_work.add_heading('This is a level 4 Heading', 4)
#Add standard lines (paragraphs)
	doc_work.add_paragraph('This is a new paragraph 1.')
	doc_work.add_paragraph('This is a new paragraph 2.',style = 'NoSpacing')
	doc_work.add_paragraph('This is a new paragraph 3.')
	doc_work.add_paragraph('This is a new paragraph 4.')
	doc_work.add_paragraph('This is a new paragraph 5.')

#add table header
	table = doc_work.add_table(rows=1, cols=5)
	table.style = 'TableGrid'
	hdr_cells = table.rows[0].cells
	hdr_cells[0].paragraphs[0].add_run('Head 1').bold = True
	hdr_cells[1].paragraphs[0].add_run('Head 2').bold = True
	hdr_cells[2].paragraphs[0].add_run('Head 3').bold = True
	hdr_cells[3].paragraphs[0].add_run('Head 4').bold = True
	hdr_cells[4].paragraphs[0].add_run('Head 5 and half').bold = True

#add table details 10 rows, 5 cells across
	for i in range(10):
		row_cells = table.add_row().cells
		for j in range(5):
			row_cells[j].text = 'row ' + str(i) + ', col ' + str(j)

#add a page break
	doc_work.add_page_break()

#Open a Template/Form .docx file that you created in word
	doc_tb = docx.Document(wk_path + '01detail_section.docx')

#merge  this into doc_work
	for element in doc_tb.element.body:
		doc_work.element.body.append(element)

#save it back as work
	doc_work.save(wk_path + 'new-file-name_as_work.docx')

# Part 2
#Use a template word file to create a non stnadard size page
#Open a Template/Form .docx file that you created in word 4 x 6 inch form
	doc = docx.Document(wk_path + 'template4x6.docx')

#list the main paragraph text (not the table text)
	for x in doc.paragraphs:
		print(x.text)

#Load the paragraphs into a list
	fullText = []
	for para in doc.paragraphs:
	    fullText.append(para.text)
#print the paragraph list
	print(fullText)

#overwrite the contents of paragraphs:  input is coded as %%Ti%%
#repalce as needed
	for para in doc.paragraphs:
		for i in range(0,10):
			if para.text == "%%T" + str(i) + "%%":
				para.text = "this is paragraph " + str(i)


#now look through the tables 0 & 1 and change paragraphs in each table
# %%Tab_i_j%%
	tables = doc.tables
	for table in tables:
	    for row in table.rows:
	        for cell in row.cells:
	            for para in cell.paragraphs:
	                for i in range(0,10):
	                	for j in range(0,10):
	                		if para.text == "%%Tab_" + str(i) + "_" + str(j) + "%%":
	                			para.text = "this is tab " + str(i) + " line " + str(j)

#save doc
	doc.save(wk_path + 'new-file-name.docx')

#Part 3
#new doc_out document for merge demo
	doc_out = docx.Document()
#merge work to out
	for element in doc_work.element.body:
		doc_out.element.body.append(element)

#merge doc into out. Note that somehow it fllows the layout of doc 4 x 6
#Note: The last document merged controls
	for element in doc.element.body:
		doc_out.element.body.append(element)

	doc_out.save(wk_path + 'new-file-name_final.docx')

#we just merged doc_work then doc.
#what happens if we change the order? try It!

if __name__  ==  '__main__':
	docx_demo(r'.\docs\\')
